"""
Content extractor and chunker for file-level retrieval.

This module handles:
- Reading text from various file formats
- Chunking text with configurable strategies
- Preserving structure (headings, code blocks)
"""

import re
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field
import sys

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from core.models import Chunk


@dataclass
class ChunkingConfig:
    """Configuration for text chunking."""
    
    # Target chunk size in tokens (approximate: 1 token ≈ 4 chars)
    chunk_size_chars: int = 2000  # ~500 tokens
    
    # Overlap between chunks in chars
    overlap_chars: int = 200  # ~50 tokens
    
    # Minimum chunk size
    min_chunk_size: int = 500
    
    # Preserve paragraph boundaries
    respect_paragraphs: bool = True
    
    # For code: try to preserve function/class boundaries
    respect_code_blocks: bool = True


class ContentExtractor:
    """Extracts and chunks content from files."""
    
    def __init__(self, config: Optional[ChunkingConfig] = None):
        self.config = config or ChunkingConfig()
    
    def extract_text(self, file_path: str) -> Optional[str]:
        """
        Extract text content from a file.
        
        Supports: txt, md, py, js, json, yaml, csv, log, etc.
        Also supports PDF and DOCX when optional dependencies are installed.
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == '.pdf':
            return self._extract_pdf_text(path)

        if ext == '.docx':
            return self._extract_docx_text(path)
        
        try:
            # Try reading as UTF-8 first
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Handle special formats
            if ext == '.json':
                # Could add JSON formatting/parsing here
                return content
            elif ext in {'.yaml', '.yml'}:
                return content
            elif ext == '.csv':
                return content
            else:
                return content
                
        except UnicodeDecodeError:
            # Try with latin-1 as fallback
            try:
                with open(path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception:
                return None
        except Exception:
            return None

    def _extract_pdf_text(self, path: Path) -> Optional[str]:
        """Extract text from a PDF file using pypdf."""
        try:
            from pypdf import PdfReader
        except Exception:
            return None

        try:
            reader = PdfReader(str(path))
            pages = []
            for page in reader.pages:
                page_text = page.extract_text() or ""
                if page_text:
                    pages.append(page_text)
            if not pages:
                return None
            return "\n\n".join(pages)
        except Exception:
            return None

    def _extract_docx_text(self, path: Path) -> Optional[str]:
        """Extract text from a DOCX file using python-docx."""
        try:
            from docx import Document
        except Exception:
            return None

        try:
            doc = Document(str(path))
            parts = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    parts.append(text)

            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))

            if not parts:
                return None
            return "\n\n".join(parts)
        except Exception:
            return None
    
    def chunk_text(self, text: str, file_path: str, is_code: bool = False) -> list[Chunk]:
        """
        Split text into chunks following the configured strategy.
        
        Args:
            text: Full text content
            file_path: Source file path
            is_code: Whether this is code (affects chunking strategy)
        
        Returns:
            List of Chunk objects
        """
        if not text or len(text.strip()) == 0:
            return []
        
        if is_code:
            return self._chunk_code(text, file_path)
        else:
            return self._chunk_prose(text, file_path)
    
    def _chunk_prose(self, text: str, file_path: str) -> list[Chunk]:
        """Chunk prose text, respecting paragraphs and headings."""
        chunks = []
        
        # Split by potential boundaries: double newlines, headings
        if self.config.respect_paragraphs:
            # Split by double newlines first
            segments = re.split(r'\n\s*\n', text)
        else:
            # Fixed-size chunks
            segments = [text[i:i + self.config.chunk_size_chars] 
                       for i in range(0, len(text), self.config.chunk_size_chars - self.config.overlap_chars)]
        
        current_chunk = ""
        current_start = 0
        
        for segment in segments:
            segment = segment.strip()
            if not segment:
                continue
            
            # If adding this segment exceeds chunk size, save current chunk
            if len(current_chunk) + len(segment) > self.config.chunk_size_chars and current_chunk:
                chunks.append(self._create_chunk(
                    text=current_chunk,
                    file_path=file_path,
                    start_offset=current_start,
                    chunk_id=f"{len(chunks)}"
                ))
                # Start new chunk with overlap
                overlap_start = max(0, len(current_chunk) - self.config.overlap_chars)
                current_chunk = current_chunk[overlap_start:]
                current_start += overlap_start
            
            current_chunk += ("\n\n" if current_chunk else "") + segment
        
        # Add final chunk
        if current_chunk and len(current_chunk) >= self.config.min_chunk_size:
            chunks.append(self._create_chunk(
                text=current_chunk,
                file_path=file_path,
                start_offset=current_start,
                chunk_id=f"{len(chunks)}"
            ))
        
        return chunks
    
    def _chunk_code(self, text: str, file_path: str) -> list[Chunk]:
        """Chunk code, trying to preserve function/class boundaries."""
        chunks = []
        
        if self.config.respect_code_blocks:
            # Try to split by function/class definitions
            # This is a simple heuristic - could be improved per language
            patterns = [
                r'(?m)^(def \w+)',      # Python functions
                r'(?m)^(class \w+)',     # Python classes
                r'(?m)^(function \w+)',  # JS functions
                r'(?m)^(\w+.*\{)',       # C-style blocks
            ]
            
            # Find all potential split points
            split_points = [0]
            for pattern in patterns:
                for match in re.finditer(pattern, text):
                    split_points.append(match.start())
            
            split_points = sorted(set(split_points))
            
            # Create segments from split points
            segments = []
            for i, start in enumerate(split_points):
                end = split_points[i + 1] if i + 1 < len(split_points) else len(text)
                segment = text[start:end].strip()
                if segment:
                    segments.append((start, segment))
            
            # Combine segments into chunks
            current_chunk = ""
            current_start = 0
            
            for start, segment in segments:
                if len(current_chunk) + len(segment) > self.config.chunk_size_chars and current_chunk:
                    chunks.append(self._create_chunk(
                        text=current_chunk,
                        file_path=file_path,
                        start_offset=current_start,
                        chunk_id=f"{len(chunks)}"
                    ))
                    current_chunk = segment
                    current_start = start
                else:
                    if not current_chunk:
                        current_start = start
                    current_chunk += ("\n\n" if current_chunk else "") + segment
            
            if current_chunk:
                chunks.append(self._create_chunk(
                    text=current_chunk,
                    file_path=file_path,
                    start_offset=current_start,
                    chunk_id=f"{len(chunks)}"
                ))
        else:
            # Fallback to fixed-size chunking
            chunks = self._chunk_prose(text, file_path)
        
        return chunks
    
    def _create_chunk(self, text: str, file_path: str, start_offset: int, chunk_id: str) -> Chunk:
        """Create a Chunk object with metadata."""
        return Chunk(
            file_path=file_path,
            chunk_id=chunk_id,
            text=text,
            start_offset=start_offset,
            end_offset=start_offset + len(text),
            metadata={
                "file": Path(file_path).name,
                "extension": Path(file_path).suffix,
            }
        )
